import pdfplumber
from docx import Document

def load_job_description(path: str) -> str:
    """Read a job description from a text or docx file."""
    if path.endswith(".txt"):
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    elif path.endswith(".docx"):
        doc = Document(path)
        return "\n".join([p.text for p in doc.paragraphs])
    else:
        raise ValueError("Unsupported file format for job description")

def load_cv(path: str) -> str:
    """Extract text from a PDF or DOCX CV."""
    if path.endswith(".pdf"):
        with pdfplumber.open(path) as pdf:
            return "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
    elif path.endswith(".docx"):
        doc = Document(path)
        return "\n".join([p.text for p in doc.paragraphs])
    else:
        raise ValueError("Unsupported file format for CV")
